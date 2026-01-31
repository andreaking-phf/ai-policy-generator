#!/usr/bin/env python3
"""
AI Policy Document Generator - DOCX Export
Generates Word documents from policy configuration JSON
"""

import json
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def create_custom_styles(doc):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Crimson Pro'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(26, 77, 46)  # Forest green
    
    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Crimson Pro'
    h2_font.size = Pt(18)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(26, 77, 46)
    
    # Heading 3 style
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'DM Sans'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(79, 121, 66)  # Sage green

def add_section_header(doc, text, level=2):
    """Add a formatted section header"""
    if level == 1:
        p = doc.add_paragraph(text, style='CustomTitle')
    elif level == 2:
        p = doc.add_paragraph(text, style='CustomH2')
    else:
        p = doc.add_paragraph(text, style='CustomH3')
    return p

def add_policy_section(doc, title, content):
    """Add a policy section with title and content"""
    # Add section title
    add_section_header(doc, title, level=3)
    
    # Add content
    lines = content.strip().split('\n')
    current_list = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('- ') or line.startswith('* '):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('## '):
            # H2 header
            add_section_header(doc, line[3:], level=2)
        elif line.startswith('### '):
            # H3 header
            add_section_header(doc, line[4:], level=3)
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.style = 'Body Text'

def generate_policy_document(config_data, output_path='ai_policy.docx'):
    """Generate a complete policy document from configuration"""
    
    doc = Document()
    create_custom_styles(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_paragraph('Artificial Intelligence Use Policy', style='CustomTitle')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    org_name = doc.add_paragraph(config_data.get('orgName', 'Public Health Department'))
    org_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org_name.runs[0].font.size = Pt(14)
    org_name.runs[0].font.bold = True
    
    date_para = doc.add_paragraph(f"Effective Date: {config_data.get('date', datetime.now().strftime('%B %d, %Y'))}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.runs[0].italic = True
    
    version = doc.add_paragraph("Policy Version: 1.0")
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version.runs[0].italic = True
    
    doc.add_page_break()
    
    # Table of contents placeholder
    add_section_header(doc, "Table of Contents", level=2)
    toc_items = [
        "I. Purpose and Scope",
        "II. Core Principles",
        "III. Policy Requirements",
        "IV. Governance and Implementation",
        "V. Contact Information"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # I. Purpose and Scope
    add_section_header(doc, "I. Purpose and Scope", level=2)
    
    purpose_text = f"""This policy establishes a framework for the responsible development, procurement, and use of artificial intelligence (AI) technologies at {config_data.get('orgName', 'this organization')}. It ensures AI systems align with our mission to protect and promote public health while upholding the highest standards of ethics, equity, privacy, and transparency."""
    doc.add_paragraph(purpose_text)
    
    scope_text = f"""This policy applies to all staff, contractors, and vendors involved in the use, development, or procurement of AI technologies on behalf of {config_data.get('orgName', 'this organization')}."""
    doc.add_paragraph(scope_text)
    
    doc.add_paragraph("It covers all AI systems including but not limited to:")
    ai_types = [
        "Generative AI tools (e.g., ChatGPT, Claude, Gemini)",
        "Machine learning models for data analysis and prediction",
        "Natural language processing systems",
        "Computer vision and image analysis tools",
        "Decision support and automation systems"
    ]
    for ai_type in ai_types:
        doc.add_paragraph(ai_type, style='List Bullet')
    
    # II. Core Principles
    add_section_header(doc, "II. Core Principles", level=2)
    
    principles_intro = f"{config_data.get('orgName', 'This organization')} commits to the following principles in all AI activities:"
    doc.add_paragraph(principles_intro)
    
    principles = [
        ("Public Health First", "AI serves public health goals and community well-being"),
        ("Equity and Justice", "AI promotes health equity and does not perpetuate discrimination"),
        ("Transparency", "AI use is open, explainable, and accountable"),
        ("Privacy Protection", "AI systems safeguard sensitive health information"),
        ("Human Oversight", "Human judgment remains central to all decisions"),
        ("Continuous Learning", "AI practices evolve with technology and evidence")
    ]
    
    for principle, description in principles:
        p = doc.add_paragraph()
        p.add_run(f"{principle}: ").bold = True
        p.add_run(description)
        p.style = 'List Bullet'
    
    doc.add_page_break()
    
    # III. Policy Requirements
    add_section_header(doc, "III. Policy Requirements", level=2)
    
    note = doc.add_paragraph("Requirements are prioritized based on organizational assessment and adapted for ")
    note.add_run(config_data.get('jurisdictionType', 'local')).italic = True
    note.add_run(" context.")
    
    # Add priority sections based on rankings
    if 'rankings' in config_data:
        rankings = config_data['rankings']
        
        # High priority (top 4)
        add_section_header(doc, "High Priority Requirements", level=3)
        for i, item in enumerate(rankings[:4]):
            add_policy_section(doc, f"{chr(65+i)}. {item['name']}", get_policy_content(item['name'], 'high', config_data.get('intensity', 'standard')))
        
        # Standard priority (5-8)
        add_section_header(doc, "Standard Requirements", level=3)
        for i, item in enumerate(rankings[4:8]):
            add_policy_section(doc, f"{chr(65+4+i)}. {item['name']}", get_policy_content(item['name'], 'medium', config_data.get('intensity', 'standard')))
    
    doc.add_page_break()
    
    # IV. Governance
    add_section_header(doc, "IV. Governance and Implementation", level=2)
    add_governance_section(doc, config_data)
    
    # V. Contact Information
    add_section_header(doc, "V. Contact Information", level=2)
    contact_text = f"""Questions about this policy should be directed to:

AI Policy Coordinator
{config_data.get('orgName', 'Public Health Department')}
[Contact information to be added]"""
    doc.add_paragraph(contact_text)
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run("""This policy was generated using the AI Policy Priority Builder tool
Based on evidence-based research from Kansas Health Institute, CDC, and other authoritative sources
Document should be reviewed by legal counsel and adapted to local requirements""")
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(79, 121, 66)
    
    # Save document
    doc.save(output_path)
    print(f"Policy document generated: {output_path}")
    return output_path

def get_policy_content(element_name, priority, intensity):
    """Get policy content for a specific element"""
    # Simplified - in production this would pull from full template database
    templates = {
        "Data Privacy & Security": {
            "high": """
### Purpose
Protect protected health information (PHI) and sensitive data in AI use.

### Requirements
- No PHI/PII in external AI systems
- Encryption required for all AI data processing
- Regular security audits
- Staff training on data privacy
- Compliance with HIPAA and applicable regulations
""",
            "medium": """
### Requirements
- Protect sensitive data when using AI
- Follow HIPAA requirements
- Staff training on data privacy
"""
        }
    }
    
    if element_name in templates:
        return templates[element_name].get(priority, templates[element_name].get("medium", "No template available"))
    else:
        return f"### {element_name}\n\nPolicy requirements to be developed based on organizational needs."

def add_governance_section(doc, config):
    """Add governance section to document"""
    
    add_section_header(doc, "A. Roles and Responsibilities", level=3)
    
    roles = [
        ("AI Governance Committee", "Oversees AI policy implementation and approves high-risk AI use cases"),
        ("Privacy Officer", "Reviews AI systems for data privacy compliance"),
        ("Equity Officer", "Assesses AI systems for bias and equity impacts"),
        ("IT Security", "Ensures AI systems meet cybersecurity requirements"),
        ("Department Heads", "Accountable for AI use within their units"),
        ("All Staff", "Required to comply with this policy")
    ]
    
    for role, responsibility in roles:
        p = doc.add_paragraph()
        p.add_run(f"{role}: ").bold = True
        p.add_run(responsibility)
        p.style = 'List Bullet'
    
    add_section_header(doc, "B. Policy Review and Updates", level=3)
    doc.add_paragraph("This policy will be reviewed annually and updated as needed to reflect:")
    
    review_items = [
        "Changes in AI technology and capabilities",
        "New legal or regulatory requirements",
        "Lessons learned from AI implementation",
        "Stakeholder feedback and concerns"
    ]
    for item in review_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_section_header(doc, "C. Compliance and Enforcement", level=3)
    doc.add_paragraph("Violations of this policy may result in:")
    
    consequences = [
        "Revocation of AI system access",
        "Mandatory retraining",
        "Disciplinary action up to and including termination",
        "Legal action where applicable"
    ]
    for consequence in consequences:
        doc.add_paragraph(consequence, style='List Bullet')

if __name__ == "__main__":
    # Example usage
    if len(sys.argv) > 1:
        # Load config from JSON file
        with open(sys.argv[1], 'r') as f:
            config = json.load(f)
    else:
        # Default configuration
        config = {
            "orgName": "Sample Public Health Department",
            "jurisdictionType": "local",
            "orgSize": "medium",
            "intensity": "standard",
            "date": datetime.now().strftime("%B %d, %Y"),
            "rankings": [
                {"name": "Data Privacy & Security", "score": 10},
                {"name": "Bias, Equity & Discrimination", "score": 9},
                {"name": "Transparency & Explainability", "score": 8},
                {"name": "Human Oversight & Accountability", "score": 8}
            ]
        }
    
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'ai_policy.docx'
    generate_policy_document(config, output_file)
