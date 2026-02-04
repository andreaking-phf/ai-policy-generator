"""
AI Policy Generator - Flask Application
A fully functional app for health departments to create AI use policies
"""
import os
import io
import json
import re
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
from models import db, PrioritizationSession, GeneratedPolicy, PolicyElement, init_policy_elements
from policy_templates import POLICY_TEMPLATES, get_template

# Try to import python-docx for DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-key-change-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///policy_data.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize database
db.init_app(app)

with app.app_context():
    db.create_all()
    init_policy_elements(db)


# =============================================================================
# Page Routes
# =============================================================================

@app.route('/')
def dashboard():
    """Main dashboard showing overview and recent activity"""
    recent_prioritizations = PrioritizationSession.query.order_by(
        PrioritizationSession.created_at.desc()
    ).limit(5).all()
    recent_policies = GeneratedPolicy.query.order_by(
        GeneratedPolicy.created_at.desc()
    ).limit(5).all()

    stats = {
        'total_prioritizations': PrioritizationSession.query.count(),
        'total_policies': GeneratedPolicy.query.count()
    }

    return render_template('dashboard.html',
                         recent_prioritizations=recent_prioritizations,
                         recent_policies=recent_policies,
                         stats=stats)


@app.route('/prioritize')
def prioritize():
    """Prioritization tool page"""
    elements = PolicyElement.query.order_by(PolicyElement.id).all()
    elements_list = [e.to_dict() for e in elements]
    return render_template('prioritize.html', elements=elements_list)


@app.route('/generate')
def generate():
    """Policy generator page"""
    prioritizations = PrioritizationSession.query.order_by(
        PrioritizationSession.created_at.desc()
    ).all()
    return render_template('generate.html', prioritizations=prioritizations)


@app.route('/history')
def history():
    """History of all sessions and policies"""
    prioritizations = PrioritizationSession.query.order_by(
        PrioritizationSession.created_at.desc()
    ).all()
    policies = GeneratedPolicy.query.order_by(
        GeneratedPolicy.created_at.desc()
    ).all()
    return render_template('history.html',
                         prioritizations=prioritizations,
                         policies=policies)


@app.route('/policy/<int:policy_id>')
def view_policy(policy_id):
    """View a specific generated policy"""
    policy = GeneratedPolicy.query.get_or_404(policy_id)
    return render_template('view_policy.html', policy=policy)


# =============================================================================
# API Routes - Prioritization
# =============================================================================

@app.route('/api/elements', methods=['GET'])
def get_elements():
    """Get all policy elements"""
    elements = PolicyElement.query.order_by(PolicyElement.id).all()
    return jsonify([e.to_dict() for e in elements])


@app.route('/api/prioritizations', methods=['GET'])
def get_prioritizations():
    """Get all prioritization sessions"""
    sessions = PrioritizationSession.query.order_by(
        PrioritizationSession.created_at.desc()
    ).all()
    return jsonify([s.to_dict() for s in sessions])


@app.route('/api/prioritizations', methods=['POST'])
def create_prioritization():
    """Save a new prioritization session"""
    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    session = PrioritizationSession(
        name=data.get('name', f"Session {datetime.now().strftime('%Y-%m-%d %H:%M')}"),
        method=data.get('method', 'weighted'),
        rankings_json=json.dumps(data.get('rankings', [])),
        raw_data_json=json.dumps(data.get('rawData', {}))
    )

    db.session.add(session)
    db.session.commit()

    return jsonify(session.to_dict()), 201


@app.route('/api/prioritizations/<int:session_id>', methods=['GET'])
def get_prioritization(session_id):
    """Get a specific prioritization session"""
    session = PrioritizationSession.query.get_or_404(session_id)
    result = session.to_dict()
    result['rankings'] = json.loads(session.rankings_json)
    result['rawData'] = json.loads(session.raw_data_json) if session.raw_data_json else {}
    return jsonify(result)


@app.route('/api/prioritizations/<int:session_id>', methods=['DELETE'])
def delete_prioritization(session_id):
    """Delete a prioritization session"""
    session = PrioritizationSession.query.get_or_404(session_id)
    db.session.delete(session)
    db.session.commit()
    return jsonify({'message': 'Deleted successfully'}), 200


# =============================================================================
# API Routes - Policy Generation
# =============================================================================

@app.route('/api/policies', methods=['GET'])
def get_policies():
    """Get all generated policies"""
    policies = GeneratedPolicy.query.order_by(
        GeneratedPolicy.created_at.desc()
    ).all()
    return jsonify([p.to_dict() for p in policies])


@app.route('/api/policies', methods=['POST'])
def create_policy():
    """Generate and save a new policy"""
    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    # Required fields
    org_name = data.get('orgName')
    if not org_name:
        return jsonify({'error': 'Organization name is required'}), 400

    rankings = data.get('rankings', [])
    if not rankings:
        return jsonify({'error': 'Rankings are required'}), 400

    # Build policy document
    config = {
        'orgName': org_name,
        'jurisdictionType': data.get('jurisdictionType', 'local'),
        'orgSize': data.get('orgSize', 'medium'),
        'intensity': data.get('intensity', 'standard'),
        'additionalContext': data.get('additionalContext', ''),
        'date': datetime.now().strftime('%B %d, %Y'),
        'rankings': rankings
    }

    policy_html = build_policy_html(config)
    policy_markdown = build_policy_markdown(config)

    policy = GeneratedPolicy(
        name=data.get('name', f"Policy for {org_name}"),
        org_name=org_name,
        jurisdiction_type=config['jurisdictionType'],
        org_size=config['orgSize'],
        intensity=config['intensity'],
        additional_context=config['additionalContext'],
        policy_html=policy_html,
        policy_markdown=policy_markdown,
        config_json=json.dumps(config),
        prioritization_id=data.get('prioritizationId')
    )

    db.session.add(policy)
    db.session.commit()

    result = policy.to_dict()
    result['policy_html'] = policy_html
    result['policy_markdown'] = policy_markdown

    return jsonify(result), 201


@app.route('/api/policies/<int:policy_id>', methods=['GET'])
def get_policy(policy_id):
    """Get a specific policy"""
    policy = GeneratedPolicy.query.get_or_404(policy_id)
    return jsonify(policy.to_dict())


@app.route('/api/policies/<int:policy_id>', methods=['DELETE'])
def delete_policy(policy_id):
    """Delete a policy"""
    policy = GeneratedPolicy.query.get_or_404(policy_id)
    db.session.delete(policy)
    db.session.commit()
    return jsonify({'message': 'Deleted successfully'}), 200


@app.route('/api/generate-preview', methods=['POST'])
def generate_preview():
    """Generate policy preview without saving"""
    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    config = {
        'orgName': data.get('orgName', 'Organization Name'),
        'jurisdictionType': data.get('jurisdictionType', 'local'),
        'orgSize': data.get('orgSize', 'medium'),
        'intensity': data.get('intensity', 'standard'),
        'additionalContext': data.get('additionalContext', ''),
        'date': datetime.now().strftime('%B %d, %Y'),
        'rankings': data.get('rankings', []),
        'references': data.get('references', [])
    }

    policy_html = build_policy_html(config)
    policy_markdown = build_policy_markdown(config)

    return jsonify({
        'html': policy_html,
        'markdown': policy_markdown
    })


@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    """Generate Word document (.docx) version of the policy"""
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'Word document generation is not available. Install python-docx: pip install python-docx'}), 500

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    config = {
        'orgName': data.get('orgName', 'Organization Name'),
        'jurisdictionType': data.get('jurisdictionType', 'local'),
        'orgSize': data.get('orgSize', 'medium'),
        'intensity': data.get('intensity', 'standard'),
        'additionalContext': data.get('additionalContext', ''),
        'date': datetime.now().strftime('%B %d, %Y'),
        'rankings': data.get('rankings', []),
        'references': data.get('references', [])
    }

    try:
        doc = build_policy_docx(config)

        # Save to BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'ai-policy-{datetime.now().strftime("%Y%m%d")}.docx'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat-refine', methods=['POST'])
def chat_refine():
    """Handle chat-based policy refinement requests"""
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    message = data.get('message', '')
    current_policy = data.get('currentPolicy', '')
    config = data.get('config', {})

    # Simple rule-based responses for common refinement requests
    # In a production system, this would integrate with an LLM API
    response = process_refinement_request(message, current_policy, config)

    return jsonify(response)


# =============================================================================
# Helper Functions
# =============================================================================

def build_policy_html(config):
    """Build HTML policy document from config"""
    org_name = config['orgName']
    date = config['date']
    intensity = config['intensity']
    jurisdiction = config['jurisdictionType']
    rankings = config['rankings']
    references = config.get('references', [])

    html = f"""
<h1>Artificial Intelligence Use Policy</h1>
<div class="doc-header">
    <p>{org_name}</p>
    <p>Effective Date: {date}</p>
    <p>Policy Version: 1.0</p>
</div>

<hr>

<h2>I. Purpose and Scope</h2>
<div class="policy-section">
    <h3>1. Purpose</h3>
    <p>This policy establishes a framework for the responsible development, procurement, and use of artificial intelligence (AI) technologies at {org_name}. It ensures AI systems align with our mission to protect and promote public health while upholding the highest standards of ethics, equity, privacy, and transparency.</p>

    <h3>2. Scope</h3>
    <p>This policy applies to all staff, contractors, and vendors involved in the use, development, or procurement of AI technologies on behalf of {org_name}. It covers all AI systems including but not limited to:</p>
    <ul>
        <li>Generative AI tools (e.g., ChatGPT, Claude, Gemini)</li>
        <li>Machine learning models for data analysis and prediction</li>
        <li>Natural language processing systems</li>
        <li>Computer vision and image analysis tools</li>
        <li>Decision support and automation systems</li>
    </ul>
</div>

<h2>II. Core Principles</h2>
<div class="policy-section">
    <p>{org_name} commits to the following principles in all AI activities:</p>
    <ul>
        <li>Public Health First: AI serves public health goals and community well-being</li>
        <li>Equity and Justice: AI promotes health equity and does not perpetuate discrimination</li>
        <li>Transparency: AI use is open, explainable, and accountable</li>
        <li>Privacy Protection: AI systems safeguard sensitive health information</li>
        <li>Human Oversight: Human judgment remains central to all decisions</li>
        <li>Continuous Learning: AI practices evolve with technology and evidence</li>
    </ul>
</div>

<hr>

<h2>III. Policy Requirements</h2>
<p>Requirements are prioritized based on organizational assessment and adapted for {jurisdiction} context.</p>
"""

    # Split rankings into priority tiers
    top_priorities = rankings[:4] if len(rankings) >= 4 else rankings
    medium_priorities = rankings[4:8] if len(rankings) >= 8 else rankings[4:]
    lower_priorities = rankings[8:] if len(rankings) > 8 else []

    # High priority sections
    html += '<h3 class="priority-high">High Priority Requirements</h3>'
    for i, item in enumerate(top_priorities):
        template = get_template(item['name'], 'high' if intensity == 'comprehensive' else 'medium' if intensity == 'standard' else 'low')
        if template:
            html += f"""
<div class="policy-section">
    <h3>{chr(65 + i)}. {template['title']}</h3>
    {markdown_to_html(template['content'])}
</div>
"""

    # Standard priority sections
    if medium_priorities:
        html += '<h3 class="priority-medium">Standard Requirements</h3>'
        for i, item in enumerate(medium_priorities):
            template = get_template(item['name'], 'medium' if intensity == 'comprehensive' else 'low')
            if template:
                html += f"""
<div class="policy-section">
    <h3>{chr(65 + len(top_priorities) + i)}. {template['title']}</h3>
    {markdown_to_html(template['content'])}
</div>
"""

    # Additional considerations (comprehensive only)
    if intensity == 'comprehensive' and lower_priorities:
        html += '<h3 class="priority-low">Additional Considerations</h3>'
        for i, item in enumerate(lower_priorities):
            template = get_template(item['name'], 'low')
            if template:
                html += f"""
<div class="policy-section">
    <h3>{chr(65 + len(top_priorities) + len(medium_priorities) + i)}. {template['title']}</h3>
    {markdown_to_html(template['content'])}
</div>
"""

    # Governance section
    html += """
<hr>

<h2>IV. Governance and Implementation</h2>
<div class="policy-section">
    <h3>A. Roles and Responsibilities</h3>
    <ul>
        <li><strong>AI Governance Committee:</strong> Oversees AI policy implementation and approves high-risk AI use cases</li>
        <li><strong>Privacy Officer:</strong> Reviews AI systems for data privacy compliance</li>
        <li><strong>Equity Officer:</strong> Assesses AI systems for bias and equity impacts</li>
        <li><strong>IT Security:</strong> Ensures AI systems meet cybersecurity requirements</li>
        <li><strong>Department Heads:</strong> Accountable for AI use within their units</li>
        <li><strong>All Staff:</strong> Required to comply with this policy</li>
    </ul>

    <h3>B. Policy Review and Updates</h3>
    <p>This policy will be reviewed annually and updated as needed to reflect:</p>
    <ul>
        <li>Changes in AI technology and capabilities</li>
        <li>New legal or regulatory requirements</li>
        <li>Lessons learned from AI implementation</li>
        <li>Stakeholder feedback and concerns</li>
    </ul>

    <h3>C. Compliance and Enforcement</h3>
    <p>Violations of this policy may result in:</p>
    <ul>
        <li>Revocation of AI system access</li>
        <li>Mandatory retraining</li>
        <li>Disciplinary action up to and including termination</li>
        <li>Legal action where applicable</li>
    </ul>
</div>
"""

    html += f"""
<h2>V. Contact Information</h2>
<div class="policy-section">
    <p>Questions about this policy should be directed to:</p>
    <p>AI Policy Coordinator<br>
    {org_name}<br>
    [Contact information to be added]</p>
</div>
"""

    # Add references section if there are any
    if references:
        html += """
<h2>VI. References and Related Policies</h2>
<div class="policy-section references-section">
    <p>This policy should be read in conjunction with the following documents:</p>
    <ul>
"""
        for ref in references:
            if ref.get('url'):
                html += f'        <li><a href="{ref["url"]}" target="_blank">{ref["title"]}</a></li>\n'
            else:
                html += f'        <li>{ref["title"]}</li>\n'
        html += """    </ul>
</div>
"""

    html += """
<hr>

<p class="footer-note">
    This policy was generated using the AI Policy Priority Builder tool.
    Based on evidence-based research from Kansas Health Institute, CDC, and other authoritative sources.
    Document should be reviewed by legal counsel and adapted to local requirements.
</p>
"""

    return html


def build_policy_markdown(config):
    """Build Markdown policy document from config"""
    org_name = config['orgName']
    date = config['date']
    intensity = config['intensity']
    jurisdiction = config['jurisdictionType']
    rankings = config['rankings']
    references = config.get('references', [])

    md = f"""# ARTIFICIAL INTELLIGENCE USE POLICY

{org_name}

Effective Date: {date}

Policy Version: 1.0

---

## I. Purpose and Scope

### 1. Purpose

This policy establishes a framework for the responsible development, procurement, and use of artificial intelligence (AI) technologies at {org_name}. It ensures AI systems align with our mission to protect and promote public health while upholding the highest standards of ethics, equity, privacy, and transparency.

### 2. Scope

This policy applies to all staff, contractors, and vendors involved in the use, development, or procurement of AI technologies on behalf of {org_name}. It covers all AI systems including but not limited to:

- Generative AI tools (e.g., ChatGPT, Claude, Gemini)
- Machine learning models for data analysis and prediction
- Natural language processing systems
- Computer vision and image analysis tools
- Decision support and automation systems

## II. Core Principles

{org_name} commits to the following principles in all AI activities:

- Public Health First: AI serves public health goals and community well-being
- Equity and Justice: AI promotes health equity and does not perpetuate discrimination
- Transparency: AI use is open, explainable, and accountable
- Privacy Protection: AI systems safeguard sensitive health information
- Human Oversight: Human judgment remains central to all decisions
- Continuous Learning: AI practices evolve with technology and evidence

---

## III. Policy Requirements

Requirements are prioritized based on organizational assessment and adapted for {jurisdiction} context.

"""

    # Split rankings into priority tiers
    top_priorities = rankings[:4] if len(rankings) >= 4 else rankings
    medium_priorities = rankings[4:8] if len(rankings) >= 8 else rankings[4:]
    lower_priorities = rankings[8:] if len(rankings) > 8 else []

    # High priority sections
    md += "### High Priority Requirements\n\n"
    for i, item in enumerate(top_priorities):
        template = get_template(item['name'], 'high' if intensity == 'comprehensive' else 'medium' if intensity == 'standard' else 'low')
        if template:
            md += f"#### {chr(65 + i)}. {template['title']}\n\n{template['content']}\n\n"

    # Standard priority sections
    if medium_priorities:
        md += "### Standard Requirements\n\n"
        for i, item in enumerate(medium_priorities):
            template = get_template(item['name'], 'medium' if intensity == 'comprehensive' else 'low')
            if template:
                md += f"#### {chr(65 + len(top_priorities) + i)}. {template['title']}\n\n{template['content']}\n\n"

    # Additional considerations
    if intensity == 'comprehensive' and lower_priorities:
        md += "### Additional Considerations\n\n"
        for i, item in enumerate(lower_priorities):
            template = get_template(item['name'], 'low')
            if template:
                md += f"#### {chr(65 + len(top_priorities) + len(medium_priorities) + i)}. {template['title']}\n\n{template['content']}\n\n"

    md += f"""---

## IV. Governance and Implementation

### A. Roles and Responsibilities

- AI Governance Committee: Oversees AI policy implementation and approves high-risk AI use cases
- Privacy Officer: Reviews AI systems for data privacy compliance
- Equity Officer: Assesses AI systems for bias and equity impacts
- IT Security: Ensures AI systems meet cybersecurity requirements
- Department Heads: Accountable for AI use within their units
- All Staff: Required to comply with this policy

### B. Policy Review and Updates

This policy will be reviewed annually and updated as needed to reflect:

- Changes in AI technology and capabilities
- New legal or regulatory requirements
- Lessons learned from AI implementation
- Stakeholder feedback and concerns

### C. Compliance and Enforcement

Violations of this policy may result in:

- Revocation of AI system access
- Mandatory retraining
- Disciplinary action up to and including termination
- Legal action where applicable

## V. Contact Information

Questions about this policy should be directed to:

AI Policy Coordinator
{org_name}
[Contact information to be added]

"""

    # Add references section if there are any
    if references:
        md += """## VI. References and Related Policies

This policy should be read in conjunction with the following documents:

"""
        for ref in references:
            if ref.get('url'):
                md += f"- [{ref['title']}]({ref['url']})\n"
            else:
                md += f"- {ref['title']}\n"
        md += "\n"

    md += """---

This policy was generated using the AI Policy Priority Builder tool.
Based on evidence-based research from Kansas Health Institute, CDC, and other authoritative sources.
Document should be reviewed by legal counsel and adapted to local requirements.
"""

    return md


def markdown_to_html(md_content):
    """Simple markdown to HTML converter"""
    html = md_content

    # Headers
    html = html.replace('### ', '<h4>').replace('\n\n<h4>', '</h4>\n\n<h4>')
    html = html.replace('## ', '<h3>').replace('\n\n<h3>', '</h3>\n\n<h3>')

    # Lists
    lines = html.split('\n')
    result = []
    in_list = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('- '):
            if not in_list:
                result.append('<ul>')
                in_list = True
            result.append(f'<li>{stripped[2:]}</li>')
        else:
            if in_list:
                result.append('</ul>')
                in_list = False
            if stripped:
                if not stripped.startswith('<'):
                    result.append(f'<p>{stripped}</p>')
                else:
                    result.append(stripped)

    if in_list:
        result.append('</ul>')

    html = '\n'.join(result)

    # Bold - only for actual emphasis, not for labels
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

    return html


def build_policy_docx(config):
    """Build Word document (.docx) from config"""
    doc = Document()

    org_name = config['orgName']
    date = config['date']
    intensity = config['intensity']
    jurisdiction = config['jurisdictionType']
    rankings = config['rankings']
    references = config.get('references', [])

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('ARTIFICIAL INTELLIGENCE USE POLICY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header info
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.add_run(f'{org_name}\n').bold = False
    header_para.add_run(f'Effective Date: {date}\n')
    header_para.add_run('Policy Version: 1.0')

    doc.add_paragraph()

    # I. Purpose and Scope
    doc.add_heading('I. Purpose and Scope', level=1)

    doc.add_heading('1. Purpose', level=2)
    doc.add_paragraph(
        f'This policy establishes a framework for the responsible development, procurement, '
        f'and use of artificial intelligence (AI) technologies at {org_name}. It ensures AI '
        f'systems align with our mission to protect and promote public health while upholding '
        f'the highest standards of ethics, equity, privacy, and transparency.'
    )

    doc.add_heading('2. Scope', level=2)
    doc.add_paragraph(
        f'This policy applies to all staff, contractors, and vendors involved in the use, '
        f'development, or procurement of AI technologies on behalf of {org_name}. It covers '
        f'all AI systems including but not limited to:'
    )

    scope_items = [
        'Generative AI tools (e.g., ChatGPT, Claude, Gemini)',
        'Machine learning models for data analysis and prediction',
        'Natural language processing systems',
        'Computer vision and image analysis tools',
        'Decision support and automation systems'
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')

    # II. Core Principles
    doc.add_heading('II. Core Principles', level=1)
    doc.add_paragraph(f'{org_name} commits to the following principles in all AI activities:')

    principles = [
        'Public Health First: AI serves public health goals and community well-being',
        'Equity and Justice: AI promotes health equity and does not perpetuate discrimination',
        'Transparency: AI use is open, explainable, and accountable',
        'Privacy Protection: AI systems safeguard sensitive health information',
        'Human Oversight: Human judgment remains central to all decisions',
        'Continuous Learning: AI practices evolve with technology and evidence'
    ]
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')

    # III. Policy Requirements
    doc.add_heading('III. Policy Requirements', level=1)
    doc.add_paragraph(
        f'Requirements are prioritized based on organizational assessment and adapted for '
        f'{jurisdiction} context.'
    )

    # Process rankings by priority tier
    top_priorities = rankings[:4] if len(rankings) >= 4 else rankings
    medium_priorities = rankings[4:8] if len(rankings) >= 8 else rankings[4:]
    lower_priorities = rankings[8:] if len(rankings) > 8 else []

    # High Priority
    if top_priorities:
        doc.add_heading('High Priority Requirements', level=2)
        for i, item in enumerate(top_priorities):
            template = get_template(item['name'], 'high' if intensity == 'comprehensive' else 'medium' if intensity == 'standard' else 'low')
            if template:
                doc.add_heading(f'{chr(65 + i)}. {template["title"]}', level=3)
                add_markdown_to_docx(doc, template['content'])

    # Standard Priority
    if medium_priorities:
        doc.add_heading('Standard Requirements', level=2)
        for i, item in enumerate(medium_priorities):
            template = get_template(item['name'], 'medium' if intensity == 'comprehensive' else 'low')
            if template:
                doc.add_heading(f'{chr(65 + len(top_priorities) + i)}. {template["title"]}', level=3)
                add_markdown_to_docx(doc, template['content'])

    # Additional Considerations
    if intensity == 'comprehensive' and lower_priorities:
        doc.add_heading('Additional Considerations', level=2)
        for i, item in enumerate(lower_priorities):
            template = get_template(item['name'], 'low')
            if template:
                doc.add_heading(f'{chr(65 + len(top_priorities) + len(medium_priorities) + i)}. {template["title"]}', level=3)
                add_markdown_to_docx(doc, template['content'])

    # IV. Governance
    doc.add_heading('IV. Governance and Implementation', level=1)

    doc.add_heading('A. Roles and Responsibilities', level=2)
    roles = [
        'AI Governance Committee: Oversees AI policy implementation and approves high-risk AI use cases',
        'Privacy Officer: Reviews AI systems for data privacy compliance',
        'Equity Officer: Assesses AI systems for bias and equity impacts',
        'IT Security: Ensures AI systems meet cybersecurity requirements',
        'Department Heads: Accountable for AI use within their units',
        'All Staff: Required to comply with this policy'
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')

    doc.add_heading('B. Policy Review and Updates', level=2)
    doc.add_paragraph('This policy will be reviewed annually and updated as needed to reflect:')
    review_items = [
        'Changes in AI technology and capabilities',
        'New legal or regulatory requirements',
        'Lessons learned from AI implementation',
        'Stakeholder feedback and concerns'
    ]
    for item in review_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('C. Compliance and Enforcement', level=2)
    doc.add_paragraph('Violations of this policy may result in:')
    enforcement_items = [
        'Revocation of AI system access',
        'Mandatory retraining',
        'Disciplinary action up to and including termination',
        'Legal action where applicable'
    ]
    for item in enforcement_items:
        doc.add_paragraph(item, style='List Bullet')

    # V. Contact Information
    doc.add_heading('V. Contact Information', level=1)
    doc.add_paragraph('Questions about this policy should be directed to:')
    doc.add_paragraph(f'AI Policy Coordinator\n{org_name}\n[Contact information to be added]')

    # VI. References (if any)
    if references:
        doc.add_heading('VI. References and Related Policies', level=1)
        doc.add_paragraph('This policy should be read in conjunction with the following documents:')
        for ref in references:
            doc.add_paragraph(ref['title'], style='List Bullet')

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        'This policy was generated using the AI Policy Priority Builder tool. '
        'Based on evidence-based research from Kansas Health Institute, CDC, and other authoritative sources. '
        'Document should be reviewed by legal counsel and adapted to local requirements.'
    ).italic = True

    return doc


def add_markdown_to_docx(doc, content):
    """Add markdown content to Word document"""
    lines = content.strip().split('\n')
    in_list = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check for headers
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=4)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=3)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith('- '):
            # List item
            text = stripped[2:]
            # Remove bold markers for cleaner output
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        else:
            # Regular paragraph - remove bold markers
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            doc.add_paragraph(text)


def process_refinement_request(message, current_policy, config):
    """Process a chat refinement request and return a response"""
    message_lower = message.lower()

    # Common refinement patterns
    if 'simplify' in message_lower or 'simpler' in message_lower:
        return {
            'response': (
                "I understand you'd like to simplify the policy. In a full implementation, "
                "I would analyze the specific section you mentioned and provide a simplified version. "
                "For now, consider these tips for simplification:\n\n"
                "1. Replace technical jargon with plain language\n"
                "2. Shorten sentences to 20 words or fewer\n"
                "3. Use bullet points instead of long paragraphs\n"
                "4. Remove redundant phrases\n\n"
                "Would you like me to focus on a specific section?"
            ),
            'updatedPolicy': None,
            'updatedHtml': None
        }

    elif 'detail' in message_lower or 'more about' in message_lower or 'expand' in message_lower:
        return {
            'response': (
                "You'd like more detail added. In a full implementation, I would expand the "
                "relevant section with additional specifics. Consider adding:\n\n"
                "1. Specific procedures and workflows\n"
                "2. Examples of compliant vs. non-compliant behavior\n"
                "3. Timelines and deadlines\n"
                "4. Responsible parties for each requirement\n\n"
                "Which section would you like me to expand?"
            ),
            'updatedPolicy': None,
            'updatedHtml': None
        }

    elif 'tone' in message_lower or 'formal' in message_lower or 'informal' in message_lower:
        return {
            'response': (
                "I can help adjust the tone of the policy. Options include:\n\n"
                "- More formal: Use passive voice, legal terminology, 'shall' instead of 'must'\n"
                "- Less formal: Use active voice, conversational language, 'should' and 'will'\n"
                "- Technical: Include more regulatory citations and technical standards\n"
                "- Accessible: Plain language for general staff understanding\n\n"
                "Which direction would you prefer?"
            ),
            'updatedPolicy': None,
            'updatedHtml': None
        }

    elif 'remove' in message_lower or 'delete' in message_lower:
        return {
            'response': (
                "I can help remove sections. To do this, please specify:\n\n"
                "1. The exact section name or letter (e.g., 'Section A' or 'Data Privacy')\n"
                "2. Or describe what content you'd like removed\n\n"
                "Note: Removing core elements may leave gaps in your AI governance. "
                "I'll let you know if a removal might create compliance issues."
            ),
            'updatedPolicy': None,
            'updatedHtml': None
        }

    elif 'add' in message_lower or 'include' in message_lower:
        return {
            'response': (
                "I can help add new content. Please specify:\n\n"
                "1. What topic or requirement you'd like to add\n"
                "2. Where in the document it should go\n"
                "3. How detailed it should be\n\n"
                "I can also suggest additions based on your organization type and "
                "common gaps in AI policies."
            ),
            'updatedPolicy': None,
            'updatedHtml': None
        }

    else:
        return {
            'response': (
                "I'm here to help refine your policy. I can:\n\n"
                "- Simplify language in specific sections\n"
                "- Add more detail to requirements\n"
                "- Adjust the tone (more/less formal)\n"
                "- Remove or add sections\n"
                "- Clarify specific terms or requirements\n\n"
                "What would you like to change? Please be specific about which section "
                "or aspect of the policy you'd like me to work on."
            ),
            'updatedPolicy': None,
            'updatedHtml': None
        }


# =============================================================================
# Main Entry Point
# =============================================================================

if __name__ == '__main__':
    app.run(debug=True, port=5000)
